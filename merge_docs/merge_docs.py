#!/usr/bin/env python
# -*- coding: utf-8 -*-
import argparse
import copy
import io
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def natural_key(text):
    return [int(part) if part.isdigit() else part.lower() for part in re.split(r'(\d+)', text)]


def collect_docx_files(folder, recursive):
    files = []
    for root, dirs, names in os.walk(folder):
        for name in names:
            if name.lower().endswith('.docx'):
                files.append(os.path.join(root, name))
        if not recursive:
            break
    return sorted(files, key=lambda p: natural_key(os.path.basename(p)))


def collect_doc_files(folder, recursive):
    files = []
    for root, dirs, names in os.walk(folder):
        for name in names:
            if name.lower().endswith('.doc'):
                files.append(os.path.join(root, name))
        if not recursive:
            break
    return sorted(files, key=lambda p: natural_key(os.path.basename(p)))


def merge_styles(source, target):
    src_root = source.styles.element
    dst_root = target.styles.element
    existing = set()
    for style_el in dst_root.findall(qn('w:style')):
        sid = style_el.get(qn('w:styleId'))
        if sid:
            existing.add(sid)
    for style_el in src_root.findall(qn('w:style')):
        sid = style_el.get(qn('w:styleId'))
        if not sid or sid in existing:
            continue
        dst_root.append(copy.deepcopy(style_el))
        existing.add(sid)


def merge_numbering(source, target):
    num_map = {}
    try:
        src_part = source.part.numbering_part
        dst_part = target.part.numbering_part
    except Exception:
        return num_map
    src_el = src_part.element
    dst_el = dst_part.element

    used_abstract_ids = set()
    used_num_ids = set()
    for ab in dst_el.findall(qn('w:abstractNum')):
        val = ab.get(qn('w:abstractNumId'))
        if val is not None:
            used_abstract_ids.add(int(val))
    for num in dst_el.findall(qn('w:num')):
        val = num.get(qn('w:numId'))
        if val is not None:
            used_num_ids.add(int(val))

    abstract_map = {}
    for ab in src_el.findall(qn('w:abstractNum')):
        src_id = int(ab.get(qn('w:abstractNumId')))
        if src_id in abstract_map:
            continue
        new_id = src_id
        if new_id in used_abstract_ids:
            new_id = max(used_abstract_ids) + 1
        abstract_map[src_id] = new_id
        used_abstract_ids.add(new_id)
        new_ab = copy.deepcopy(ab)
        new_ab.set(qn('w:abstractNumId'), str(new_id))
        dst_el.append(new_ab)

    for num in src_el.findall(qn('w:num')):
        src_id = int(num.get(qn('w:numId')))
        new_id = src_id
        if new_id in used_num_ids:
            new_id = max(used_num_ids) + 1
        num_map[src_id] = new_id
        used_num_ids.add(new_id)
        new_num = copy.deepcopy(num)
        for abs_el in new_num.findall(qn('w:abstractNumId')):
            old_abs = int(abs_el.get(qn('w:val')))
            abs_el.set(qn('w:val'), str(abstract_map[old_abs]))
        new_num.set(qn('w:numId'), str(new_id))
        dst_el.append(new_num)

    return num_map


def remap_relationships(source, target, element):
    for el in element.iter():
        for attr_name in ('r:embed', 'r:id', 'r:link'):
            rel_id = el.get(qn(attr_name))
            if not rel_id:
                continue
            rel = source.part.rels.get(rel_id)
            if rel is None:
                continue
            if rel.reltype == RT.IMAGE:
                stream = io.BytesIO(rel.target_part.blob)
                new_rel_id = target.part.get_or_add_image(stream)[0]
                el.set(qn(attr_name), new_rel_id)
            elif rel.reltype == RT.HYPERLINK and rel.is_external:
                new_rel_id = target.part.relate_to(rel.target_ref, RT.HYPERLINK, is_external=True)
                el.set(qn(attr_name), new_rel_id)


def remap_numbering(element, num_map):
    if not num_map:
        return
    for pPr in element.findall(qn('w:pPr')):
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            continue
        num_id_el = numPr.find(qn('w:numId'))
        if num_id_el is None:
            continue
        val = num_id_el.get(qn('w:val'))
        if val and val.isdigit() and int(val) in num_map:
            num_id_el.set(qn('w:val'), str(num_map[int(val)]))


def append_to_body(dst_body, element):
    sect_prs = [c for c in dst_body if c.tag == qn('w:sectPr')]
    if sect_prs:
        sect_prs[-1].addprevious(element)
    else:
        dst_body.append(element)


def append_source_body(source, target, num_map):
    src_body = source.element.body
    dst_body = target.element.body
    p_count = 0
    t_count = 0
    if src_body.find(qn('w:altChunk')) is not None:
        print('    警告：该文档包含 altChunk 嵌入内容，可能无法完整合并。')
    for child in list(src_body):
        if child.tag == qn('w:sectPr'):
            continue
        new_element = copy.deepcopy(child)
        try:
            remap_relationships(source, target, new_element)
        except Exception as exc:
            print('  警告：处理图片/链接失败（%s）' % exc)
        remap_numbering(new_element, num_map)
        if new_element.tag == qn('w:p'):
            p_count += 1
        elif new_element.tag == qn('w:tbl'):
            t_count += 1
        append_to_body(dst_body, new_element)
    return p_count, t_count


def extract_header_text(source):
    texts = []
    try:
        for section in source.sections:
            for p in section.header.paragraphs:
                if p.text.strip():
                    texts.append(p.text.strip())
    except Exception:
        pass
    return texts


def add_document_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_toc(doc, levels):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-%d" \\h \\z \\u' % levels
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = '（在 Word 中全选并按 F9，或右键选择“更新域”以生成目录）'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def merge_documents(files, output_path, toc_levels):
    target = Document()

    target.add_heading('目录', level=1)
    add_toc(target, toc_levels)
    target.add_page_break()

    errors = []
    for index, path in enumerate(files):
        try:
            source = Document(path)
            merge_styles(source, target)
            num_map = merge_numbering(source, target)
            if index > 0:
                target.add_page_break()
            add_document_title(target, os.path.splitext(os.path.basename(path))[0])
            p_count, t_count = append_source_body(source, target, num_map)
            img_count = len(source.inline_shapes)
            print('  合并：%s（段落 %d，表格 %d，图片 %d）' % (path, p_count, t_count, img_count))
            if p_count == 0 and t_count == 0 and img_count == 0:
                header_texts = extract_header_text(source)
                if header_texts:
                    for text in header_texts:
                        target.add_paragraph(text)
                    print('    注意：该文档正文为空，已将页眉内容（%d 段）复制到正文。' % len(header_texts))
                else:
                    print('    注意：该文档正文为空，仅添加了文件名分隔。')
        except Exception as exc:
            errors.append((path, str(exc)))
            print('  跳过（处理失败）：%s（%s）' % (path, exc))

    target.save(output_path)
    return errors


def main():
    parser = argparse.ArgumentParser(
        description='将文件夹内所有 Word(.docx) 文档合并为一个文档，保留原文结构并生成可检索目录')
    parser.add_argument('folder', nargs='?', default='.', help='存放 Word 文档的文件夹（默认：当前目录）')
    parser.add_argument('-o', '--output', help='输出文件路径（默认：merged_日期时间.docx，放在 folder 下）')
    parser.add_argument('-r', '--recursive', action='store_true', help='递归扫描子文件夹中的文档')
    parser.add_argument('-l', '--toc-levels', type=int, default=3, help='目录包含的标题级别（默认 3）')
    args = parser.parse_args()

    folder = os.path.abspath(args.folder)
    if not os.path.isdir(folder):
        print('文件夹不存在：%s' % folder)
        sys.exit(1)

    files = collect_docx_files(folder, args.recursive)
    legacy_docs = collect_doc_files(folder, args.recursive)

    if legacy_docs:
        print('发现旧版 .doc 文档（无法直接处理，请先用 Word 另存为 .docx）：')
        for f in legacy_docs:
            print('  - %s' % f)

    if not files:
        print('在 %s 中未找到任何 .docx 文档。' % folder)
        sys.exit(1)

    output = args.output or 'merged_%s.docx' % datetime.now().strftime('%Y%m%d_%H%M%S')
    if not os.path.isabs(output):
        output = os.path.join(folder, output)
    output_abs = os.path.abspath(output)
    files = [f for f in files if os.path.abspath(f).lower() != output_abs.lower()]

    print('共找到 %d 个 .docx 文档：' % len(files))
    for f in files:
        print('  - %s' % f)

    errors = merge_documents(files, output, args.toc_levels)

    print('\n合并完成：%s' % output)
    if errors:
        print('以下文档合并失败：')
        for path, msg in errors:
            print('  - %s（%s）' % (path, msg))
    print('提示：在 Word 中打开后，全选并按 F9（或右键选择“更新域”）刷新目录；打印前请再次更新页码。')


if __name__ == '__main__':
    main()
